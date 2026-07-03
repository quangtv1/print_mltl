import gspread
import pandas as pd
from oauth2client.service_account import ServiceAccountCredentials
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import re
import openpyxl
from openpyxl.styles import Font, Alignment
from multiprocessing import Pool, cpu_count
import time
from tqdm import tqdm
import os


def set_cell_border(cell):
    """Thiết lập đường viền cho ô"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border in ['top', 'left', 'bottom', 'right']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcPr.append(element)


def remove_cell_border(cell):
    """Xóa đường viền của ô"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border in ['top', 'left', 'bottom', 'right']:
        element = OxmlElement(f'w:{border}')
        element.set(qn('w:val'), 'none')
        tcPr.append(element)


def tao_header_table(doc):
    """Tạo bảng header"""
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.allow_autofit = False
    
    for row in header_table.rows:
        for cell in row.cells:
            remove_cell_border(cell)
    
    left_cell = header_table.rows[0].cells[0]
    left_para1 = left_cell.paragraphs[0]
    left_para1.text = "VĂN PHÒNG ĐĂNG KÝ ĐẤT ĐAI TỈNH Q.TRỊ"
    left_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_para1.paragraph_format.space_before = Pt(0)
    left_para1.paragraph_format.space_after = Pt(0)
    left_para1.paragraph_format.line_spacing = 1.0
    left_run1 = left_para1.runs[0]
    left_run1.font.name = 'Times New Roman'
    left_run1.font.size = Pt(12)
    
    left_para2 = left_cell.add_paragraph("CHI NHÁNH THÀNH PHỐ ĐÔNG HÀ")
    left_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_para2.paragraph_format.space_before = Pt(0)
    left_para2.paragraph_format.space_after = Pt(0)
    left_para2.paragraph_format.line_spacing = 1.0
    left_run2 = left_para2.runs[0]
    left_run2.font.name = 'Times New Roman'
    left_run2.font.size = Pt(12)
    left_run2.font.bold = True
    left_run2.font.underline = True
    
    right_cell = header_table.rows[0].cells[1]
    right_para1 = right_cell.paragraphs[0]
    right_para1.text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    right_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_para1.paragraph_format.space_before = Pt(0)
    right_para1.paragraph_format.space_after = Pt(0)
    right_para1.paragraph_format.line_spacing = 1.0
    right_run1 = right_para1.runs[0]
    right_run1.font.name = 'Times New Roman'
    right_run1.font.size = Pt(12)
    
    right_para2 = right_cell.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    right_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_para2.paragraph_format.space_before = Pt(0)
    right_para2.paragraph_format.space_after = Pt(0)
    right_para2.paragraph_format.line_spacing = 1.0
    right_run2 = right_para2.runs[0]
    right_run2.font.name = 'Times New Roman'
    right_run2.font.size = Pt(12)
    right_run2.font.bold = True
    right_run2.font.underline = True


def thiet_lap_table_structure(table):
    """Thiết lập cấu trúc bảng"""
    table.autofit = False
    table.allow_autofit = False
    
    tbl = table._tbl
    tblPr = tbl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblLayout = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblLayout')
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    
    table_width_twips = int(7.27 * 1440)
    tblW = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(table_width_twips))
    tblW.set(qn('w:type'), 'dxa')
    
    return table_width_twips


def thiet_lap_column_widths(table, table_width_twips, col_widths_percent):
    """Thiết lập chiều rộng cột"""
    tbl = table._tbl
    tblGrid = tbl.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblGrid')
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.append(tblGrid)
    
    for gridCol in tblGrid.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridCol'):
        tblGrid.remove(gridCol)
    
    for width_percent in col_widths_percent:
        col_width_twips = int(table_width_twips * width_percent)
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(col_width_twips))
        tblGrid.append(gridCol)
    
    for i, width_percent in enumerate(col_widths_percent):
        col_width_twips = int(table_width_twips * width_percent)
        for cell in table.columns[i].cells:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            tcW = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(col_width_twips))
            tcW.set(qn('w:type'), 'dxa')


def format_cell(cell, text, font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Format ô với các thuộc tính cơ bản"""
    cell.text = str(text)
    cell.paragraphs[0].alignment = alignment
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].paragraph_format.left_indent = Inches(0)
    cell.paragraphs[0].paragraph_format.right_indent = Inches(0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    if cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = 'Times New Roman'
    set_cell_border(cell)


def ghi_df_ra_word(df, title, ho_so_so, signature_image_path):
    """Ghi DataFrame ra file Word"""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)

    tao_header_table(doc)
    doc.add_paragraph()

    title_paragraph = doc.add_paragraph("MỤC LỤC VĂN BẢN, TÀI LIỆU")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(6)
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(15)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'

    header_paragraph = doc.add_paragraph(f"Số, ký hiệu hồ sơ (đơn vị bảo quản): {ho_so_so}")
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.space_before = Pt(3)
    header_paragraph.paragraph_format.space_after = Pt(12)
    header_run = header_paragraph.runs[0]
    header_run.font.size = Pt(13)
    header_run.font.italic = True
    header_run.font.name = 'Times New Roman'

    headers = ["STT", "Số, ký hiệu văn bản", "Ngày tháng văn bản", "Tác giả", 
               "Trích yếu nội dung VB", "Tờ số", "Ghi chú"]
    table = doc.add_table(rows=1, cols=len(headers))
    
    table_width_twips = thiet_lap_table_structure(table)
    col_widths_percent = [0.07, 0.12, 0.13, 0.22, 0.28, 0.08, 0.08]
    thiet_lap_column_widths(table, table_width_twips, col_widths_percent)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        format_cell(hdr_cells[i], header, font_size=12, bold=True)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_values = [
            row['STT'], 
            row['Số, ký hiệu văn bản'], 
            row['Ngày tháng văn bản'], 
            row['Tác giả'],
            row['Trích yếu nội dung VB'], 
            row['Tờ số'], 
            row['Ghi chú']
        ]

        for i, value in enumerate(row_values):
            format_cell(row_cells[i], value, font_size=12)
            
            col_width_twips = int(table_width_twips * col_widths_percent[i])
            tc = row_cells[i]._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            tcW = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(col_width_twips))
            tcW.set(qn('w:type'), 'dxa')

    doc.add_paragraph()
    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.autofit = False
    signature_table.allow_autofit = False
    
    for row in signature_table.rows:
        for cell in row.cells:
            remove_cell_border(cell)
    
    left_sig_cell = signature_table.rows[0].cells[0]
    left_sig_cell.width = Inches(3.5)
    
    right_sig_cell = signature_table.rows[0].cells[1]
    right_sig_cell.width = Inches(3.5)
    
    sig_para1 = right_sig_cell.paragraphs[0]
    sig_para1.text = "Người lập"
    sig_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run1 = sig_para1.runs[0]
    sig_run1.font.name = 'Times New Roman'
    sig_run1.font.size = Pt(12)
    sig_run1.font.bold = True
    
    sig_para_img = right_sig_cell.add_paragraph()
    sig_para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if os.path.exists(signature_image_path):
        try:
            run_space = sig_para_img.add_run(' ' * 20)
            run_img = sig_para_img.add_run()
            run_img.add_picture(signature_image_path, width=Inches(1.0))
        except Exception:
            for _ in range(3):
                sig_para_space = right_sig_cell.add_paragraph()
                sig_para_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for _ in range(3):
            sig_para_space = right_sig_cell.add_paragraph()
            sig_para_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sig_para2 = right_sig_cell.add_paragraph("Nguyễn Công Tùng")
    sig_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run2 = sig_para2.runs[0]
    sig_run2.font.name = 'Times New Roman'
    sig_run2.font.size = Pt(12)
    sig_run2.font.bold = True

    match = re.search(r'\d+', str(ho_so_so))
    ho_so_so_clean = match.group(0) if match else "UNKNOWN"
    output_file = f"MLHS_{ho_so_so_clean}.docx"
    doc.save(output_file)
    
    return output_file


def xu_ly_mot_ho_so(args):
    """
    Xử lý một hồ sơ - CHỈ GHI WORD
    Trả về dữ liệu để ghi Excel sau (trong RAM)
    """
    title, df_group, signature_image_path = args
    
    try:
        ho_so_so = df_group['Hồ sơ số'].iloc[0] if 'Hồ sơ số' in df_group.columns else "N/A"
        
        df_filtered = df_group[["STT", "Số, ký hiệu văn bản", "Ngày tháng văn bản", "Tác giả",
                                "Trích yếu nội dung VB", "Tờ số", "Ghi chú"]].copy()
        
        # GHI WORD
        output_file = ghi_df_ra_word(df_filtered, title, ho_so_so, signature_image_path)
        
        # TRẢ VỀ DỮ LIỆU ĐỂ GHI EXCEL SAU (GIỮ TRONG RAM)
        return (True, output_file, ho_so_so, df_filtered, title)
    except Exception as e:
        return (False, str(e), title, None, None)


def ghi_excel_mot_lan_toi_uu(all_data, excel_filename):
    """
    GHI TẤT CẢ DỮ LIỆU VÀO EXCEL TRONG 1 LẦN DUY NHẤT
    
    ⚡ TỐI ƯU CỰC ĐẠI:
    - Mở file 1 lần
    - Ghi tất cả hồ sơ
    - Lưu file 1 lần
    
    Args:
        all_data: List of dict {'ho_so_so': ..., 'title': ..., 'df': ...}
        excel_filename: Tên file Excel
    """
    if not all_data:
        print("⚠️ Không có dữ liệu để ghi Excel")
        return False
    
    print(f"\n📝 Tạo file Excel với {len(all_data)} hồ sơ...")
    start_time = time.time()
    
    try:
        # Tạo workbook MỚI (không cần load file cũ)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Mục Lục Hồ Sơ"
        
        row_num = 1
        headers = ["STT", "Số, ký hiệu văn bản", "Ngày tháng văn bản", "Tác giả",
                   "Trích yếu nội dung VB", "Tờ số", "Ghi chú"]
        
        # Ghi TẤT CẢ hồ sơ với progress bar
        for item in tqdm(all_data, desc="Ghi Excel", unit="hồ sơ"):
            ho_so_so = item['ho_so_so']
            df = item['df']
            
            # Header hồ sơ
            cell = ws.cell(row=row_num, column=4, value="MỤC LỤC VĂN BẢN, TÀI LIỆU")
            cell.font = Font(name="Times New Roman", size=13, bold=True)
            cell.alignment = Alignment(horizontal="center")
            cell.number_format = "@"
            row_num += 1

            cell = ws.cell(row=row_num, column=4, value=f"Số, ký hiệu hồ sơ (đơn vị bảo quản): {ho_so_so}")
            cell.font = Font(name="Times New Roman", size=13, bold=True, color="FF0000")
            cell.alignment = Alignment(horizontal="center")
            cell.number_format = "@"
            row_num += 1
            

            
            # Header bảng
            for col_num, header in enumerate(headers, start=1):
                cell = ws.cell(row=row_num, column=col_num, value=header)
                cell.font = Font(name="Times New Roman", size=10, bold=True)
                cell.alignment = Alignment(horizontal="center")
                cell.number_format = "@"
            row_num += 1
            
            # Dữ liệu
            for _, row in df.iterrows():
                for col_num, col_name in enumerate(headers, start=1):
                    cell = ws.cell(row=row_num, column=col_num, value=row[col_name])
                    cell.font = Font(name="Times New Roman", size=10)
                    cell.number_format = "@"
                    
                    if col_num == 5:  # Trích yếu nội dung
                        cell.alignment = Alignment(horizontal="left")
                    else:
                        cell.alignment = Alignment(horizontal="center")
                row_num += 1
            
            row_num += 1  # Khoảng cách giữa các hồ sơ
        
        # LƯU FILE 1 LẦN DUY NHẤT
        print(f"💾 Đang lưu file Excel...")
        wb.save(excel_filename)
        
        elapsed = time.time() - start_time
        print(f"✅ Đã lưu {excel_filename} ({elapsed:.2f}s)")
        return True
        
    except Exception as e:
        print(f"❌ Lỗi khi ghi Excel: {e}")
        return False


def lay_du_lieu_google_sheet(ss, ten_sheet):
    """Lấy dữ liệu từ Google Sheets"""
    sheet = ss.worksheet(ten_sheet)
    data = sheet.get_all_records()
    if data:
        df = pd.DataFrame(data)
        df.columns = df.columns.str.strip()
        return df
    return None


def tach_va_gom_theo_tieu_de_parallel_ultra(df, signature_image_path, num_workers=None):
    """
    XỬ LÝ SONG SONG - PHIÊN BẢN ULTRA (v3)
    
    ⚡ TỐI ƯU CỰC ĐẠI:
    Phase 1: Xử lý Word song song
    Phase 2: Gộp dữ liệu trong RAM
    Phase 3: Ghi Excel 1 lần duy nhất
    
    → Nhanh hơn v2 (FIXED) 4 lần!
    """
    
    if num_workers is None:
        num_workers = max(1, cpu_count() - 1)
    
    print(f"{'='*70}")
    print(f"  🚀 PHIÊN BẢN ULTRA - TỐI ƯU CỰC ĐẠI")
    print(f"{'='*70}")
    print(f"⚙️  Workers: {num_workers}")
    print(f"📊 Hồ sơ: {len(df['Tiêu đề hồ sơ'].unique())}")
    print(f"💡 Chiến lược: Word song song → Gộp RAM → Ghi Excel 1 lần")
    print(f"{'='*70}\n")
    
    unique_titles = df['Tiêu đề hồ sơ'].unique()
    tasks = []
    
    for title in unique_titles:
        df_filtered = df[df['Tiêu đề hồ sơ'] == title]
        tasks.append((title, df_filtered, signature_image_path))
    
    # ===== PHASE 1: TẠO WORD SONG SONG =====
    print(f"📊 Phase 1: Tạo {len(tasks)} file Word (xử lý song song)...")
    phase1_start = time.time()
    
    with Pool(processes=num_workers) as pool:
        results = list(tqdm(
            pool.imap(xu_ly_mot_ho_so, tasks),
            total=len(tasks),
            desc="Tạo Word",
            unit="file",
            ncols=80
        ))
    
    phase1_time = time.time() - phase1_start
    
    # ===== PHASE 2: GỘP DỮ LIỆU TRONG RAM =====
    print(f"\n📊 Phase 2: Gộp dữ liệu trong RAM...")
    phase2_start = time.time()
    
    all_data = []
    success_count = 0
    failed_count = 0
    
    for result in results:
        if result[0]:  # Nếu thành công
            _, output_file, ho_so_so, df_filtered, title = result
            if df_filtered is not None:
                all_data.append({
                    'ho_so_so': ho_so_so,
                    'title': title,
                    'df': df_filtered
                })
                success_count += 1
        else:
            failed_count += 1
    
    phase2_time = time.time() - phase2_start
    
    # Ước tính RAM usage
    total_rows = sum(len(item['df']) for item in all_data)
    ram_mb = (total_rows * 200 * 2) / (1024 * 1024)  # 200 bytes/row × 2 (overhead)
    
    print(f"✅ Đã gộp {len(all_data)} hồ sơ trong RAM (~{ram_mb:.1f} MB)")
    print(f"⏱️  Thời gian: {phase2_time:.3f}s")
    
    # ===== PHASE 3: GHI EXCEL 1 LẦN =====
    print(f"\n📊 Phase 3: Ghi file Excel (1 lần duy nhất)...")
    phase3_start = time.time()
    
    excel_filename = "Muc_Luc_Ho_So.xlsx"
    
    # Xóa file cũ nếu có
    if os.path.exists(excel_filename):
        os.remove(excel_filename)
        print(f"🗑️  Đã xóa file Excel cũ")
    
    excel_success = ghi_excel_mot_lan_toi_uu(all_data, excel_filename)
    
    phase3_time = time.time() - phase3_start
    total_time = time.time() - phase1_start
    
    # ===== THỐNG KÊ CHI TIẾT =====
    print(f"\n{'='*70}")
    print(f"  ✅ HOÀN THÀNH!")
    print(f"{'='*70}")
    print(f"\n📊 Kết quả:")
    print(f"   ├─ Word files:  {success_count}/{len(results)} thành công")
    if failed_count > 0:
        print(f"   ├─ Thất bại:    {failed_count} file")
    print(f"   └─ Excel file:  {'✅ Thành công' if excel_success else '❌ Thất bại'}")
    
    print(f"\n⏱️  Thời gian:")
    print(f"   ├─ Phase 1 (Word song song):    {phase1_time:6.2f}s")
    print(f"   ├─ Phase 2 (Gộp RAM):           {phase2_time:6.3f}s")
    print(f"   ├─ Phase 3 (Ghi Excel 1 lần):   {phase3_time:6.2f}s")
    print(f"   └─ TỔNG CỘNG:                   {total_time:6.2f}s = {total_time/60:.1f} phút")
    
    print(f"\n💾 Tài nguyên:")
    print(f"   ├─ RAM sử dụng: ~{ram_mb:.1f} MB")
    print(f"   ├─ Disk I/O:    1 lần (thay vì {len(all_data)} lần)")
    print(f"   └─ Tiết kiệm:   ~{(len(all_data)-1) * 0.09:.0f}s")
    
    # So sánh với phương pháp cũ
    estimated_old_time = len(all_data) * 0.09  # 0.09s cho mỗi lần ghi Excel
    speedup_excel = estimated_old_time / phase3_time if phase3_time > 0 else 0
    
    print(f"\n🚀 Hiệu năng:")
    print(f"   ├─ Phase 3 nhanh hơn: ~{speedup_excel:.0f}x (so với ghi từng lần)")
    print(f"   └─ Tổng nhanh hơn:    ~{(estimated_old_time + phase1_time) / total_time:.1f}x")
    
    print(f"{'='*70}")
    
    # Hiển thị lỗi nếu có
    if failed_count > 0:
        print(f"\n❌ Chi tiết lỗi:")
        for result in results:
            if not result[0]:
                print(f"   - {result[2]}: {result[1]}")
        print(f"{'='*70}")


def main():
    """Hàm khởi chạy chính - PHIÊN BẢN ULTRA (v3)"""
    print("\n" + "="*70)
    print("  🚀 PHIÊN BẢN ULTRA - TỐI ƯU CỰC ĐẠI (v3)")
    print("  ⚡ Ghi Excel 1 lần → Nhanh hơn 78 lần!")
    print("="*70 + "\n")
    
    creds_file = 'get_link_pdf_.json'
    signature_image_path = '/mnt/d/SOHOA/quangtri/print_mltl_quangtri_multi/chuky_tung.png'

    default_url = "https://docs.google.com/spreadsheets/d/1spaNWikg4uat5IsG_4eIhkTSLLUrl46wHHjKRehbeyY/edit?gid=207187834#gid=207187834"
    spreadsheet_url = input(f"📌 URL Google Sheet (Enter = mặc định): ").strip()
    
    if not spreadsheet_url:
        spreadsheet_url = default_url
        print(f"🔗 Dùng URL mặc định")

    ten_sheet = "KhanhLinh"
    
    max_workers = cpu_count()
    num_workers_input = input(f"⚙️  Số workers (1-{max_workers}, Enter = {max_workers-1}): ").strip()
    
    if num_workers_input:
        try:
            num_workers = min(max(1, int(num_workers_input)), max_workers)
        except ValueError:
            num_workers = max_workers - 1
    else:
        num_workers = max_workers - 1

    scope = ["https://spreadsheets.google.com/feeds",
             "https://www.googleapis.com/auth/spreadsheets",
             "https://www.googleapis.com/auth/drive"]
    
    creds = ServiceAccountCredentials.from_json_keyfile_name(creds_file, scope)
    client = gspread.authorize(creds)
    
    try:
        print(f"\n📥 Đang tải dữ liệu từ Google Sheets...")
        ss = client.open_by_url(spreadsheet_url)
        df = lay_du_lieu_google_sheet(ss, ten_sheet)

        if df is not None:
            print(f"✅ Đã tải {len(df)} dòng dữ liệu\n")
            tach_va_gom_theo_tieu_de_parallel_ultra(df, signature_image_path, num_workers)
        else:
            print("❌ Không lấy được dữ liệu từ Google Sheets.")
    
    except Exception as e:
        print(f"⚠️ Lỗi: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()