"""
Sinh template mặc định cho app "Tạo Mục Lục Hồ Sơ".

Tái tạo CHÍNH XÁC layout đang hardcode trong print_mltl_parallel.py
(header cơ quan, bảng 7 cột, khối chữ ký), nhưng thay các giá trị cứng
bằng placeholder docxtpl:
  - Biến cấp tài liệu:  {{ co_quan_dong1 }}, {{ ho_so_so }}, {{ nguoi_ky }}...
  - Vòng lặp hàng bảng: {%tr for r in rows %} ... {%tr endfor %} với {{ r.stt }}...
  - Ảnh chữ ký:         {{ anh_chu_ky }} (docxtpl InlineImage lúc render)
Đồng thời chèn footer động "Trang {PAGE}/{NUMPAGES}" canh phải.

Kết quả:
  styles/van-phong-dat-dai/template.docx
  styles/van-phong-dat-dai/style.json

Chạy: python3 build_default_template.py
Chỉ cần python-docx (không cần docxtpl để dựng template).
"""

import json
import os

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


# ----------------------------------------------------------------------------
# Helper OOXML (giữ nguyên logic từ print_mltl_parallel.py — phần dựng khung)
# ----------------------------------------------------------------------------

FONT_NAME = "Times New Roman"


def set_cell_border(cell):
    """Đường viền single cho ô (khớp bản cũ: val=single, sz=4)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for border in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{border}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcPr.append(el)


def remove_cell_border(cell):
    """Bỏ viền ô (dùng cho bảng header & bảng chữ ký)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for border in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{border}")
        el.set(qn("w:val"), "none")
        tcPr.append(el)


def style_run(run, size, bold=False, italic=False, underline=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline


def add_centered_line(cell, text, size, bold=False, underline=False, first=False):
    """Thêm 1 dòng canh giữa vào ô, khoảng cách 0, line spacing 1.0."""
    para = cell.paragraphs[0] if first else cell.add_paragraph()
    para.text = text
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    style_run(para.runs[0], size, bold=bold, underline=underline)
    return para


def tao_header_table(doc):
    """Bảng header 2 cột không viền: trái = cơ quan (biến), phải = quốc hiệu."""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.allow_autofit = False
    for row in t.rows:
        for cell in row.cells:
            remove_cell_border(cell)

    left = t.rows[0].cells[0]
    add_centered_line(left, "{{ co_quan_dong1 }}", 12, first=True)
    add_centered_line(left, "{{ co_quan_dong2 }}", 12, bold=True, underline=True)

    right = t.rows[0].cells[1]
    add_centered_line(right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 12, first=True)
    add_centered_line(right, "Độc lập - Tự do - Hạnh phúc", 12, bold=True, underline=True)


def thiet_lap_table_structure(table):
    """Bảng layout fixed, tổng rộng 7.27 inch (khớp bản cũ)."""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    tblPr = tbl.find(f".//{ns}tblPr")
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblLayout = tblPr.find(f".//{ns}tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    table_width_twips = int(7.27 * 1440)
    tblW = tblPr.find(f".//{ns}tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(table_width_twips))
    tblW.set(qn("w:type"), "dxa")
    return table_width_twips


def thiet_lap_column_widths(table, table_width_twips, col_widths_percent):
    """Đặt tblGrid + tcW cho từng cột theo phần trăm."""
    tbl = table._tbl
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    tblGrid = tbl.find(f".//{ns}tblGrid")
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.append(tblGrid)
    for gridCol in tblGrid.findall(f".//{ns}gridCol"):
        tblGrid.remove(gridCol)

    for pct in col_widths_percent:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(table_width_twips * pct)))
        tblGrid.append(gridCol)

    for i, pct in enumerate(col_widths_percent):
        w = int(table_width_twips * pct)
        for cell in table.columns[i].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(f".//{ns}tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")


def format_cell(cell, text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Format ô bảng dữ liệu (viền, canh, font) — khớp bản cũ."""
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = alignment
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if p.runs:
        style_run(p.runs[0], size, bold=bold)
    set_cell_border(cell)


def add_page_field(paragraph, instr, size=12):
    """Chèn field Word (PAGE / NUMPAGES) vào paragraph bằng w:fldSimple."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))  # half-points
    rPr.append(rFonts)
    rPr.append(sz)
    t = OxmlElement("w:t")
    t.text = "1"  # giá trị placeholder, Word tự tính lại
    run.append(rPr)
    run.append(t)
    fld.append(run)
    paragraph._p.append(fld)


def them_footer_so_trang(doc):
    """Footer canh phải: 'Trang {PAGE}/{NUMPAGES}' (Word tự cập nhật x/y)."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Trang ")
    style_run(r, 12)
    add_page_field(p, "PAGE")
    r2 = p.add_run("/")
    style_run(r2, 12)
    add_page_field(p, "NUMPAGES")


# ----------------------------------------------------------------------------
# Dựng template
# ----------------------------------------------------------------------------

# Cột bảng: (nhãn header, biến docxtpl trong loop, phần trăm rộng)
COLUMNS = [
    ("STT", "stt", 0.07),
    ("Số, ký hiệu văn bản", "so_ky_hieu_vb", 0.12),
    ("Ngày tháng văn bản", "ngay_thang_vb", 0.13),
    ("Tác giả", "tac_gia", 0.22),
    ("Trích yếu nội dung VB", "trich_yeu", 0.28),
    ("Tờ số", "to_so", 0.08),
    ("Ghi chú", "ghi_chu", 0.08),
]


def build_template(path):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.5)

    # Header cơ quan / quốc hiệu
    tao_header_table(doc)
    doc.add_paragraph()

    # Tiêu đề văn bản
    title = doc.add_paragraph("{{ tieu_de }}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(6)
    style_run(title.runs[0], 15, bold=True)

    # Dòng "Số, ký hiệu hồ sơ ..."
    ho_so = doc.add_paragraph(
        "Số, ký hiệu hồ sơ (đơn vị bảo quản): {{ ho_so_so }}"
    )
    ho_so.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ho_so.paragraph_format.space_before = Pt(3)
    ho_so.paragraph_format.space_after = Pt(12)
    style_run(ho_so.runs[0], 13, italic=True)

    # Bảng: hàng header + hàng-marker {%tr for%} + hàng dữ liệu + hàng-marker {%tr endfor%}
    # docxtpl XÓA nguyên hàng chứa tag {%tr%} và biến thành lệnh jinja, nên
    # for/endfor phải ở 2 hàng marker riêng, kẹp hàng dữ liệu lặp ở giữa.
    table = doc.add_table(rows=4, cols=len(COLUMNS))
    table_width_twips = thiet_lap_table_structure(table)
    thiet_lap_column_widths(table, table_width_twips, [c[2] for c in COLUMNS])

    # Hàng 0: header (nhãn cố định)
    hdr = table.rows[0].cells
    for i, (label, _, _) in enumerate(COLUMNS):
        format_cell(hdr[i], label, size=12, bold=True)

    # Hàng 1: marker mở vòng lặp (sẽ bị docxtpl xóa)
    table.rows[1].cells[0].text = "{%tr for r in rows %}"

    # Hàng 2: hàng dữ liệu lặp thực sự
    body = table.rows[2].cells
    for i, (_, var, _) in enumerate(COLUMNS):
        align = (
            WD_ALIGN_PARAGRAPH.LEFT if var == "trich_yeu"
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        format_cell(body[i], "{{ r." + var + " }}", size=12, alignment=align)

    # Hàng 3: marker đóng vòng lặp (sẽ bị docxtpl xóa)
    table.rows[3].cells[0].text = "{%tr endfor %}"

    # Khối chữ ký (bảng 2 cột không viền, dồn phải)
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = False
    sig.allow_autofit = False
    for row in sig.rows:
        for cell in row.cells:
            remove_cell_border(cell)
    sig.rows[0].cells[0].width = Inches(3.5)
    right = sig.rows[0].cells[1]
    right.width = Inches(3.5)

    add_centered_line(right, "{{ chuc_danh_ky }}", 12, bold=True, first=True)
    # Bọc `{% if %}` để khi thiếu ảnh (anh_chu_ky=None) không render chữ "None",
    # chỉ chừa 1 dòng trống — renderer truyền InlineImage hoặc None (xem P3).
    img_para = right.add_paragraph("{% if anh_chu_ky %}{{ anh_chu_ky }}{% endif %}")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_centered_line(right, "{{ nguoi_ky }}", 12, bold=True)

    # Footer số trang động
    them_footer_so_trang(doc)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def build_style_json(path):
    """Cấu hình mặc định: mapping cột→biến, cột gom nhóm, settings."""
    style = {
        "name": "Văn phòng đăng ký đất đai - Đông Hà",
        "template_file": "template.docx",
        "grouping_column": "Tiêu đề hồ sơ",
        "output_filename_pattern": "MLHS_{ho_so_so}.docx",
        "settings": {
            "co_quan_dong1": "VĂN PHÒNG ĐĂNG KÝ ĐẤT ĐAI TỈNH Q.TRỊ",
            "co_quan_dong2": "CHI NHÁNH THÀNH PHỐ ĐÔNG HÀ",
            "tieu_de": "MỤC LỤC VĂN BẢN, TÀI LIỆU",
            "chuc_danh_ky": "Người lập",
            "nguoi_ky": "Nguyễn Công Tùng",
            "anh_chu_ky_path": "chuky_tung.png",
            "font_name": FONT_NAME,
            "footer_page_number": True,
            "footer_format": "Trang {PAGE}/{NUMPAGES}",
        },
        # Biến cấp tài liệu: biến docxtpl -> tên cột trong sheet
        "document_fields": {
            "ho_so_so": "Hồ sơ số",
        },
        # Biến trong vòng lặp hàng: biến docxtpl -> tên cột trong sheet
        # (gợi ý mặc định cho auto-match; người dùng sửa được trong app)
        "row_mapping": {
            "stt": "STT",
            "so_ky_hieu_vb": "Số, ký hiệu văn bản",
            "ngay_thang_vb": "Ngày tháng văn bản",
            "tac_gia": "Tác giả",
            "trich_yeu": "Trích yếu nội dung VB",
            "to_so": "Tờ số",
            "ghi_chu": "Ghi chú",
        },
        # Định nghĩa cột bảng (nhãn + độ rộng %) để UI settings hiển thị/sửa
        "columns": [
            {"header": h, "var": v, "width": w} for (h, v, w) in COLUMNS
        ],
    }
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(style, f, ensure_ascii=False, indent=2)
    return path


if __name__ == "__main__":
    base = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "styles", "van-phong-dat-dai")
    tpl = build_template(os.path.join(base, "template.docx"))
    js = build_style_json(os.path.join(base, "style.json"))
    print(f"✅ Đã tạo template mặc định:\n   {tpl}\n   {js}")
